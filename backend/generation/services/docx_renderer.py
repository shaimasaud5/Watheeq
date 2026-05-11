# generation/services/docx_renderer.py
# ──────────────────────────────────────
# Receives structured data from the orchestrator and generates a formatted Word document.
# Supports both BRD and MOM document structures based on the predefined schemas.

import os
from datetime import date
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .prompting import key_to_label


# Colors
NAVY  = RGBColor(0x1F, 0x39, 0x64)
BLUE  = RGBColor(0x2E, 0x75, 0xB6)
GREY  = RGBColor(0x40, 0x40, 0x40)
BLACK = RGBColor(0x00, 0x00, 0x00)

# Normalize priority and status values
PRIORITY_MAP = {
    "must": "High", "must-have": "High",
    "should": "Medium", "should-have": "Medium",
    "could": "Low", "could-have": "Low",
    "high": "High", "medium": "Medium", "low": "Low",
}
STATUS_MAP = {
    "approved": "Approved", "pending": "Pending",
    "open": "Open", "closed": "Closed", "done": "Done",
    "in progress": "In Progress",
}


# ═══════════════════════════════════════════════════════════════
# # Main document builder
# ═══════════════════════════════════════════════════════════════

def build_doc_from_json(
    doc:          Document,
    ordered:      Dict[str, Any],
    doc_type:     str,
    project_name: str = "",
    cover_meta:   Dict[str, str] = None,
) -> None:
    """
    Build the complete Word document.

    Parameters:
        doc          — Empty Word document object
        ordered      — Structured and ordered data from the orchestrator
        doc_type     — "BRD" or "MOM"
        project_name — Project name used in the cover page
        cover_meta   — {date, author, version}
    """
    _set_margins(doc)
    _apply_styles(doc)

    # Page 1: Cover page
    _add_cover_page(doc, doc_type, project_name, cover_meta or {})
    doc.add_page_break()
    _add_footer(doc)

    # Render sections in the same order received from the orchestrator
    for section_name, section_value in ordered.items():
        if section_value is None:
            continue
        if isinstance(section_value, (list, dict)) and len(section_value) == 0:
            continue
        _render_section(doc, section_name, section_value)


# ═══════════════════════════════════════════════════════════════
# Render a single section
# ═══════════════════════════════════════════════════════════════

def _render_section(doc: Document, name: str, value: Any) -> None:
    """
    Render a single section based on its value type:

    str          → text paragraph
    list of str  → bullet list
    list of dict → structured item blocks
    dict         → nested subsections
    """
    _add_heading1(doc, key_to_label(name))

    if isinstance(value, str):
        _add_paragraph(doc, value)

    elif isinstance(value, list):
        if not value:
            return
        if isinstance(value[0], dict):
            for idx, item in enumerate(value, start=1):
                _render_item_block(doc, item, idx)
        else:
            for item in value:
                text = str(item).strip()
                if text:
                    _add_bullet(doc, text)

    elif isinstance(value, dict):
        for sub_name, sub_value in value.items():
            if sub_value is None or sub_value == "" or sub_value == []:
                continue
            _add_heading2(doc, key_to_label(sub_name))
            if isinstance(sub_value, str):
                _add_paragraph(doc, sub_value)
            elif isinstance(sub_value, list):
                for item in sub_value:
                    text = str(item).strip()
                    if text:
                        _add_bullet(doc, text)

    _add_divider(doc)


# ═══════════════════════════════════════════════════════════════
# Render item block (list of dictionaries)
# ═══════════════════════════════════════════════════════════════
def _render_item_block(doc: Document, item: dict, idx: int) -> None:
    """
    Render a single item from a list of dictionaries.

    Example:
        [Heading 2] Item title
                    Priority: High
                    Description: Enriched text...
    """
    # Heading priority:
    # title → risk → task → term → requirement → name → Item N
    heading_keys = ("title", "risk", "task", "term", "requirement", "name")
    heading_text = next(
        (str(item[k]).strip() for k in heading_keys if k in item and item[k]),
        f"Item {idx}"
    )
    # Shorten very long headings
    if len(heading_text) > 80:
        heading_text = heading_text[:80] + "…"

    _add_heading2(doc, heading_text)

    # Render remaining fields
    skip_as_heading = {heading_text}
    for key, val in item.items():
        if not val or str(val).strip() in ("", "None", "null"):
            continue
        if str(val).strip() in skip_as_heading:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)

        run_label = p.add_run(f"{key_to_label(key)}: ")
        run_label.bold = True
        p.add_run(_format_value(key, str(val)))

    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(6)


def _format_value(key: str, value: str) -> str:
    """Normalize priority and status values."""
    k = key.lower()
    v = value.lower().strip()
    if k == "priority":
        return PRIORITY_MAP.get(v, value.title())
    if k in ("impact", "likelihood"):
        return value.capitalize()
    if k == "status":
        return STATUS_MAP.get(v, value.title())
    return value


# ═══════════════════════════════════════════════════════════════
# Cover page
# ═══════════════════════════════════════════════════════════════

def _add_cover_page(
    doc:          Document,
    doc_type:     str,
    project_name: str,
    cover_meta:   Dict[str, str],
) -> None:
    """Build the document cover page."""
    _add_bar(doc, "1F3964", 0.6)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _add_logo(doc)

    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Document type
    titles = {"BRD": "Business Requirements Document", "MOM": "Minutes of Meeting"}
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(titles.get(doc_type, "Document").upper())
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = BLUE

    # Project name
    if project_name:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(20)
        r2 = p2.add_run(project_name)
        r2.font.name = "Calibri"
        r2.font.size = Pt(22)
        r2.font.bold = True
        r2.font.color.rgb = NAVY

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(6)
    line.paragraph_format.space_after  = Pt(14)
    _add_bottom_border(line)

    for _ in range(3):
        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Cover information table
    _add_cover_table(doc, cover_meta)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    _add_bar(doc, "2E75B6", 0.35)


def _add_cover_table(doc: Document, cover_meta: Dict[str, str]) -> None:
    """Add a borderless table for Date, Author, and Version."""
    rows_data = [
        ("Date",    cover_meta.get("date",    str(date.today()))),
        ("Author",  cover_meta.get("author",  "—")),
        ("Version", cover_meta.get("version", "1.0")),
    ]
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    _hide_borders(table)

    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        r1  = row.cells[0].paragraphs[0].add_run(label)
        r1.font.name = "Calibri"
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = GREY

        r2 = row.cells[1].paragraphs[0].add_run(str(value))
        r2.font.name = "Calibri"
        r2.font.size = Pt(10)
        r2.font.color.rgb = NAVY


def _add_logo(doc: Document) -> None:
    """Add the Watheeq logo if available; otherwise, skip silently."""
    candidates = [
        os.environ.get("WATHEEQ_LOGO_PATH", ""),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "watheeq_logo.jpg"),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "photo_5960662120447806859_x.jpg"),
    ]
    logo_path = next((p for p in candidates if p and os.path.exists(p)), None)
    if not logo_path:
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(4)
    para.add_run().add_picture(logo_path, width=Inches(1.8))


# ═══════════════════════════════════════════════════════════════
# Helper Functions for Building Document Elements
# ═══════════════════════════════════════════════════════════════

def _add_heading1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = BLUE

def _add_heading2(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = GREY

def _add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)

def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)

def _add_divider(doc: Document) -> None:
    """Thin divider line displayed after each section."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)


# ═══════════════════════════════════════════════════════════════
# Helper functions for cover page and layout
# ═══════════════════════════════════════════════════════════════

def _add_bar(doc: Document, colour: str, height_cm: float) -> None:
    """Add a horizontal colored bar using a table."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        borders.append(b)
    tblPr.append(borders)

    cell = table.rows[0].cells[0]
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  colour)
    tcPr.append(shd)

    tr   = table.rows[0]._tr
    trPr = OxmlElement("w:trPr")
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(int(height_cm * 567)))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)
    tr.insert(0, trPr)

    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after  = Pt(0)


def _hide_borders(table) -> None:
		"""Remove all table borders."""
		tbl   = table._tbl
		tblPr = tbl.find(qn("w:tblPr"))
		if tblPr is None:
			tblPr = OxmlElement("w:tblPr")
			tbl.insert(0, tblPr)
		borders = OxmlElement("w:tblBorders")
		for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
			b = OxmlElement(f"w:{name}")
			b.set(qn("w:val"), "none")
			b.set(qn("w:sz"), "0")
			b.set(qn("w:space"), "0")
			b.set(qn("w:color"), "auto")
			borders.append(b)
		tblPr.append(borders)


def _add_bottom_border(paragraph, colour: str = "2E75B6") -> None:
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "12")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), colour)
    pBdr.append(bot)
    pPr.append(pBdr)


def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)


def _apply_styles(doc: Document) -> None:
    s = doc.styles

    normal = s["Normal"]
    normal.font.name      = "Calibri"
    normal.font.size      = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after  = Pt(6)
    normal.paragraph_format.line_spacing = Pt(14)

    h1 = s["Heading 1"]
    h1.font.name      = "Calibri"
    h1.font.size      = Pt(14)
    h1.font.bold      = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before   = Pt(16)
    h1.paragraph_format.space_after    = Pt(4)
    h1.paragraph_format.keep_with_next = True

    h2 = s["Heading 2"]
    h2.font.name      = "Calibri"
    h2.font.size      = Pt(12)
    h2.font.bold      = True
    h2.font.italic    = False
    h2.font.color.rgb = GREY
    h2.paragraph_format.space_before   = Pt(10)
    h2.paragraph_format.space_after    = Pt(2)
    h2.paragraph_format.keep_with_next = True

    try:
        s["List Bullet"].font.name = "Calibri"
        s["List Bullet"].font.size = Pt(11)
    except KeyError:
        pass


def _add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    para   = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()

    def _field(run, name):
        b = OxmlElement("w:fldChar")
        b.set(qn("w:fldCharType"), "begin")
        run._r.append(b)
        i = OxmlElement("w:instrText")
        i.text = name
        run._r.append(i)
        e = OxmlElement("w:fldChar")
        e.set(qn("w:fldCharType"), "end")
        run._r.append(e)

    for text, is_field, name in [
        ("Page ", False, None), ("", True, "PAGE"),
        (" of ",  False, None), ("", True, "NUMPAGES"),
    ]:
        r = para.add_run(text)
        r.font.size = Pt(9)
        if is_field:
            _field(r, name)