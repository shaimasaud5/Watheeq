import PyPDF2
import docx
import re


def extract_text_from_pdf_file(file):
    reader = PyPDF2.PdfReader(file)
    text = ""

    for page in reader.pages:
        text += page.extract_text() or ""

    return clean_template_text(text)

def extract_text_from_docx_file(file):
    doc = docx.Document(file)

    parts = []

    for para in doc.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                parts.append(" | ".join(row_text))

    text = "\n".join(parts)
    return clean_template_text(text)

def extract_text_from_template(file):
    filename = file.name.lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf_file(file)

    elif filename.endswith(".docx"):
        return extract_text_from_docx_file(file)

    else:
        raise ValueError("Unsupported file type. Only PDF and DOCX are allowed.")

def clean_template_text(text):
    if not text:
        return text

    lines = text.split("\n")
    cleaned_lines = []

    for line in lines:
        line = line.strip()

        if not line:
            continue

        if "Download our free" in line:
            continue

        cleaned_lines.append(line)

    text = "\n".join(cleaned_lines)

    # فصل الكلمات الملصوقة بين lowercase و Uppercase
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)

    # إضافة سطر قبل العناوين المرقمة
    text = re.sub(r'(\d+\.)', r'\n\1', text)
    
    
    text = re.sub(r':\s*', ':\n', text)
    # فصل الكلمات Title Case المتتالية (مثل Name Job role Duties)
    text = re.sub(r'(?<=[a-z])\s(?=[A-Z])', '\n', text)

    # تنظيف المسافات فقط بدون كسر الأسطر
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n+', '\n', text)

    

    return text.strip()